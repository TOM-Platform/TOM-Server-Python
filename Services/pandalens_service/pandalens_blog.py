from docx import Document
from docx.shared import Inches
from Utilities import file_utility, time_utility
from Services.pandalens_service import pandalens_const


def create_blog_in_word(intro, conclusion, moments):
    # Create a new Document
    doc = Document()

    # Intro
    doc.add_heading('My Travel Blog', 0)
    doc.add_paragraph(intro)

    # Body
    doc.add_heading('Places I Visited', 0)
    for moment in moments:
        img_name = moment[pandalens_const.DB_MOMENT_FILENAME_KEY]
        img_folder = file_utility.get_path_from_project_root(pandalens_const.SEQUENCE_PATH_TO_IMAGES_FROM_PROJECT_ROOT)
        img_path = file_utility.get_path_with_filename(img_folder, img_name)
        summary = moment[pandalens_const.DB_MOMENT_SUMMARY_KEY]
        doc.add_picture(img_path, width=Inches(4))
        doc.paragraph = doc.add_paragraph("\n" + summary + "\n")

    # Conclusion
    doc.add_heading('Verdict', 0)
    doc.add_paragraph(conclusion)

    # Save the document
    pandalens_blog_root = file_utility.get_path_from_project_root(
        pandalens_const.SEQUENCE_PATH_TO_BLOGS_FROM_PROJECT_ROOT)

    current_date = time_utility.get_date_string("%Y%m%d_%H%M")
    doc.save(file_utility.get_path_with_filename(pandalens_blog_root, f'blog-{current_date}.docx'))
